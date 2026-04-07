#!/usr/bin/env python3
"""
Word 文档读取脚本

用法:
    python read-word.py <file.docx>
    python read-word.py <file.docx> --output output.txt
    python read-word.py <file.docx> --format markdown
"""

import sys
import argparse
from pathlib import Path

def read_word(file_path: str, output_format: str = "text") -> str:
    """读取 Word 文档内容"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx")
        print("运行: pip install python-docx")
        sys.exit(1)

    doc = Document(file_path)
    result = []

    if output_format == "markdown":
        result.append(f"# {Path(file_path).stem}\n")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 根据样式判断标题
        if output_format == "markdown":
            style_name = para.style.name.lower()
            if "heading 1" in style_name:
                result.append(f"\n## {text}\n")
            elif "heading 2" in style_name:
                result.append(f"\n### {text}\n")
            elif "heading 3" in style_name:
                result.append(f"\n#### {text}\n")
            else:
                result.append(text)
        else:
            result.append(text)

    # 读取表格
    if doc.tables:
        result.append("\n\n---\n## 表格内容\n")
        for i, table in enumerate(doc.tables, 1):
            result.append(f"\n### 表格 {i}\n")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if output_format == "markdown":
                    result.append("| " + " | ".join(cells) + " |")
                else:
                    result.append("  ".join(cells))

    return "\n".join(result)


def main():
    parser = argparse.ArgumentParser(description="读取 Word 文档")
    parser.add_argument("file", help="Word 文件路径 (.docx)")
    parser.add_argument("--output", "-o", help="输出文件路径")
    parser.add_argument("--format", "-f", choices=["text", "markdown"],
                        default="markdown", help="输出格式 (默认: markdown)")

    args = parser.parse_args()

    if not Path(args.file).exists():
        print(f"错误: 文件不存在: {args.file}")
        sys.exit(1)

    content = read_word(args.file, args.format)

    if args.output:
        Path(args.output).write_text(content, encoding="utf-8")
        print(f"已输出到: {args.output}")
    else:
        print(content)


if __name__ == "__main__":
    main()
