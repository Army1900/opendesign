#!/usr/bin/env python3
"""
Word 文档读取脚本

用法:
    python read-word.py <file.docx>
    python read-word.py <file.docx> --output output.txt
"""

import sys
import argparse
from pathlib import Path


def read_word(file_path: str) -> str:
    """读取 Word 文档内容"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 需要安装 python-docx")
        print("运行: pip install python-docx")
        sys.exit(1)

    doc = Document(file_path)
    result = [f"# {Path(file_path).stem}\n"]

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 根据样式判断标题级别
            style_name = para.style.name.lower()
            if "heading 1" in style_name:
                result.append(f"\n## {text}\n")
            elif "heading 2" in style_name:
                result.append(f"\n### {text}\n")
            elif "heading 3" in style_name:
                result.append(f"\n#### {text}\n")
            else:
                result.append(text)

    # 读取表格
    for table in doc.tables:
        result.append("\n\n| " + " | ".join(cell.text.strip() for cell in table.rows[0].cells) + " |")
        result.append("| " + " | ".join("---" for _ in table.rows[0].cells) + " |")
        for row in table.rows[1:]:
            result.append("| " + " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells) + " |")

    return "\n".join(result)


def main():
    parser = argparse.ArgumentParser(description="读取 Word 文档")
    parser.add_argument("file", help="Word 文件路径")
    parser.add_argument("--output", "-o", help="输出文件路径")

    args = parser.parse_args()

    if not Path(args.file).exists():
        print(f"错误: 文件不存在: {args.file}")
        sys.exit(1)

    content = read_word(args.file)

    if args.output:
        Path(args.output).write_text(content, encoding="utf-8")
        print(f"已输出到: {args.output}")
    else:
        print(content)


if __name__ == "__main__":
    main()
